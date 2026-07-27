"""
Service report generation module for equipment service reports.
"""
import os
import tempfile
import logging
from datetime import datetime

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from django.conf import settings
from django.http import HttpResponse

from .export_converter import convert_docx_to_pdf
from .models import OrganizationSettings
from .service_report_sorting import get_report_equipment_key, get_report_equipment_name, sort_service_reports

logger = logging.getLogger(__name__)


STATUS_VI_MAP = {
    'completed': 'Đã khắc phục',
    'pending': 'Chưa khắc phục',
    'temporary': 'Khắc phục tạm thời',
}


class ServiceReportGenerator:
    def __init__(self):
        self.template_dir = os.path.join(settings.BASE_DIR, 'QAID_Manager', 'templates', 'reports')
        os.makedirs(self.template_dir, exist_ok=True)

    def get_template_path(self):
        org_settings = OrganizationSettings.get_settings()
        if getattr(org_settings, 'service_report_template', None) and org_settings.service_report_template.name:
            return org_settings.service_report_template.path
        return os.path.join(self.template_dir, 'service_report_template.docx')

    def generate_doc(self, reports, date_from=None, date_to=None):
        template_path = self.get_template_path()
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            # Graceful fallback when template file is unavailable.
            doc = Document()
            doc.add_heading('Equipment Service Report', level=1)
            doc.add_paragraph('Generated without external template file.')
        now = datetime.now()
        replacements = {
            '{{REPORT_PRINTING_DATE}}': now.strftime('%d'),
            '{{REPORT_PRINTING_MONTH}}': now.strftime('%m'),
            '{{REPORT_PRINTING_YEAR}}': now.strftime('%Y'),
        }
        self._replace_text_placeholders(doc, replacements)
        self._replace_service_table_placeholder(doc, reports, date_from=date_from, date_to=date_to)
        return doc

    def _replace_text_placeholders(self, doc, replacements):
        def replace_in_paragraph(paragraph):
            for run in paragraph.runs:
                text = run.text
                for key, value in replacements.items():
                    if key in text:
                        text = text.replace(key, value)
                run.text = text

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

    def _replace_service_table_placeholder(self, doc, reports, date_from=None, date_to=None):
        placeholder = '{{SERVICE_TABLE}}'

        # Preferred path: placeholder in a table cell.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        for p in cell.paragraphs:
                            p.text = p.text.replace(placeholder, '')
                        report_table = cell.add_table(rows=1, cols=6)
                        self._fill_service_table(report_table, reports, date_from=date_from, date_to=date_to)
                        return

        # Fallback: placeholder in normal paragraph; insert table before it.
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, '')
                table = doc.add_table(rows=1, cols=6)
                self._fill_service_table(table, reports, date_from=date_from, date_to=date_to)
                paragraph._p.addprevious(table._tbl)
                return

        # If no placeholder exists, append table to end.
        table = doc.add_table(rows=1, cols=6)
        self._fill_service_table(table, reports, date_from=date_from, date_to=date_to)

    def _fill_service_table(self, table, reports, date_from=None, date_to=None):
        headers = [
            'Thiết Bị',
            'Mô tả lỗi',
            'Thời điểm ghi nhận',
            'Tình trạng',
            'Thời gian dừng máy',
            'Ghi chú',
        ]
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, title in enumerate(headers):
            table.rows[0].cells[i].text = title

        sorted_reports = sort_service_reports(reports, date_from=date_from, date_to=date_to)

        current_equipment_key = None
        group_start_row = None
        for row_index, report in enumerate(sorted_reports, start=1):
            row = table.add_row().cells
            equipment_key = get_report_equipment_key(report)
            equipment_name = get_report_equipment_name(report)

            if current_equipment_key is None:
                current_equipment_key = equipment_key
                group_start_row = row_index
                row[0].text = equipment_name
            elif equipment_key != current_equipment_key:
                self._merge_equipment_group(table, group_start_row, row_index - 1)
                current_equipment_key = equipment_key
                group_start_row = row_index
                row[0].text = equipment_name
            else:
                # Keep one visible equipment name even if a converter ignores cell merges.
                row[0].text = ''

            row[1].text = report.issue_description or ''
            row[2].text = report.date.strftime('%d/%m/%Y') if report.date else ''
            row[3].text = STATUS_VI_MAP.get(report.status, report.get_status_display())
            row[4].text = '' if report.downtime_hours is None else f"{report.downtime_hours:g} giờ"
            row[5].text = (report.notes or '').strip()

        # Merge the final contiguous equipment group in first column.
        if sorted_reports and group_start_row is not None:
            self._merge_equipment_group(table, group_start_row, len(sorted_reports))

        # Column widths: make 3rd and 5th narrower, expand 2nd and 4th.
        widths = [Inches(1.3), Inches(2.3), Inches(1.0), Inches(1.8), Inches(1.0), Inches(1.6)]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths):
                    cell.width = widths[idx]

        # Table formatting: borders, Times New Roman 13, centered, bold header.
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                self._set_cell_borders(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(13)
                        if row_idx == 0:
                            run.bold = True

    def _set_cell_borders(self, cell):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in('w:tcBorders')
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        for edge in ('top', 'left', 'bottom', 'right'):
            tag = f'w:{edge}'
            elem = borders.find(qn(tag))
            if elem is None:
                elem = OxmlElement(tag)
                borders.append(elem)
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), '8')
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), '000000')

    def _merge_equipment_group(self, table, start_row, end_row):
        """Merge first-column cells for one contiguous equipment group."""
        if start_row is None or end_row is None or end_row <= start_row:
            return
        table.cell(start_row, 0).merge(table.cell(end_row, 0))

    def save_report(self, doc, filename, output_format='pdf'):
        temp_dir = tempfile.gettempdir()
        temp_docx = os.path.join(temp_dir, f"{filename}.docx")
        temp_pdf = os.path.join(temp_dir, f"{filename}.pdf")
        doc.save(temp_docx)

        if output_format == 'docx':
            return temp_docx

        ok, errors = convert_docx_to_pdf(temp_docx, temp_pdf, allow_native_fallback=False)
        if ok and os.path.exists(temp_pdf) and os.path.getsize(temp_pdf) > 0:
            return temp_pdf
        raise RuntimeError("PDF conversion failed. " + " | ".join(errors))

    def build_response(self, reports, output_format='pdf', preview=False, filename=None, date_from=None, date_to=None):
        if filename is None:
            filename = f"service_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        output_format = (output_format or 'pdf').strip().lower()
        if output_format not in ('pdf', 'docx'):
            output_format = 'pdf'

        doc = self.generate_doc(reports, date_from=date_from, date_to=date_to)
        try:
            temp_file = self.save_report(doc, filename, output_format=output_format)
            if output_format == 'pdf' and not temp_file.endswith('.pdf'):
                raise RuntimeError("PDF export requested but non-PDF file was produced.")
        except Exception as exc:
            logger.exception("Service report export failed.")
            if output_format == 'pdf':
                return HttpResponse(
                    "PDF export failed after template-to-DOCX conversion. "
                    "Please verify Microsoft Word COM registration (or install LibreOffice) and try again. "
                    f"Technical detail: {exc}",
                    status=500,
                    content_type='text/plain; charset=utf-8',
                )
            return HttpResponse(
                f"Failed to export report: {exc}",
                status=500,
                content_type='text/plain; charset=utf-8',
            )

        if temp_file.endswith('.pdf'):
            content_type = 'application/pdf'
            ext = 'pdf'
        else:
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ext = 'docx'

        with open(temp_file, 'rb') as handle:
            response = HttpResponse(handle.read(), content_type=content_type)
            disposition = 'inline' if preview else 'attachment'
            response['Content-Disposition'] = f'{disposition}; filename="{filename}.{ext}"'

        try:
            os.remove(temp_file)
        except OSError:
            pass
        return response
