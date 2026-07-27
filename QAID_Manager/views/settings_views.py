"""
Settings Management Views.

This module contains view functions for managing application settings:
- LINACs management
- Dosimeters management
- Physics Parameters
- QA Tests configuration
- Vietnamese Holidays
- Organization Settings
- User management
- Devices and Equipment management
"""
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib import messages
from django.contrib.auth.models import User
from django.http import FileResponse, Http404
from django.views.decorators.http import require_POST
from django.db.models import Max
from ..models import Linac, Dosimeter, PhysicsParameters, QATest, VietnameseHoliday, OrganizationSettings, UserActivity, UserProfile, LinacDocument, Device, DeviceDocument, CustomTestType, CustomTest
from ..forms import (
    DosimeterForm, LinacForm, PhysicsParametersForm,
    QATestForm, VietnameseHolidayForm, OrganizationSettingsForm, UserForm, LinacDocumentForm,
    DeviceForm, DeviceDocumentForm, CustomTestTypeForm, CustomTestForm
)
from ..film_parameters_service import deduplicate_film_analysis_parameters
import os
from io import BytesIO


@login_required
def settings_home(request):
    """Settings home page showing all settings categories"""
    return render(request, 'QAID_Manager/settings_home.html')


@login_required
def settings_dosimeters(request):
    """Display list of all dosimeters (view-only for non-admin)"""
    dosimeters = Dosimeter.objects.all().order_by('name')
    is_admin = request.user.is_staff
    
    return render(request, 'QAID_Manager/settings_dosimeters.html', {
        'dosimeters': dosimeters,
        'is_admin': is_admin
    })


# LINAC Views
@login_required
def settings_linacs(request):
    """Display list of all LINACs (view-only for non-admin)"""
    linacs = Linac.objects.all().order_by('name')
    is_admin = request.user.is_staff
    
    return render(request, 'QAID_Manager/settings_linacs.html', {
        'linacs': linacs,
        'is_admin': is_admin
    })


@login_required
def linac_detail(request, pk):
    """View LINAC details"""
    linac = get_object_or_404(Linac, pk=pk)
    documents = linac.documents.all()
    is_admin = request.user.is_staff
    
    # Handle file upload
    if request.method == 'POST' and is_admin:
        form = LinacDocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.linac = linac
            document.uploaded_by = request.user
            
            # Determine file type from extension
            filename = document.file.name
            ext = filename.split('.')[-1].lower() if '.' in filename else ''
            file_type_map = {
                'pdf': 'pdf',
                'doc': 'doc',
                'docx': 'docx',
                'png': 'png',
                'jpg': 'jpg',
                'jpeg': 'jpeg',
                'tiff': 'tiff',
                'tif': 'tif',
            }
            document.file_type = file_type_map.get(ext, 'other')
            
            # Set file_name if not provided
            if not document.file_name:
                document.file_name = os.path.basename(filename)
            
            document.save()
            messages.success(request, f'File "{document.file_name}" uploaded successfully.')
            return redirect('linac_detail', pk=pk)
        else:
            messages.error(request, 'Error uploading file. Please check the form.')
    else:
        form = LinacDocumentForm()
    
    return render(request, 'QAID_Manager/linac_detail.html', {
        'linac': linac,
        'documents': documents,
        'form': form,
        'is_admin': is_admin
    })

@login_required
@staff_member_required
@require_POST
def linac_document_delete(request, pk, doc_pk):
    """Delete a LINAC document"""
    document = get_object_or_404(LinacDocument, pk=doc_pk, linac_id=pk)
    file_name = document.file_name
    document.delete()
    messages.success(request, f'File "{file_name}" deleted successfully.')
    return redirect('linac_detail', pk=pk)


@staff_member_required
def linac_create(request):
    """Create a new LINAC (admin only)"""
    if request.method == 'POST':
        form = LinacForm(request.POST)
        if form.is_valid():
            linac = form.save()
            messages.success(request, f'LINAC "{linac.name}" created successfully.')
            return redirect('settings_linacs')
    else:
        form = LinacForm()
    
    return render(request, 'QAID_Manager/linac_form.html', {
        'form': form,
        'title': 'Add New LINAC'
    })


@staff_member_required
def linac_edit(request, pk):
    """Edit an existing LINAC (admin only)"""
    linac = get_object_or_404(Linac, pk=pk)
    
    if request.method == 'POST':
        form = LinacForm(request.POST, instance=linac)
        if form.is_valid():
            linac = form.save()
            messages.success(request, f'LINAC "{linac.name}" updated successfully.')
            return redirect('settings_linacs')
    else:
        form = LinacForm(instance=linac)
    
    return render(request, 'QAID_Manager/linac_form.html', {
        'form': form,
        'linac': linac,
        'title': f'Edit LINAC: {linac.name}'
    })


@staff_member_required
def linac_delete(request, pk):
    """Delete a LINAC (admin only)"""
    linac = get_object_or_404(Linac, pk=pk)
    
    if request.method == 'POST':
        name = linac.name
        linac.delete()
        messages.success(request, f'LINAC "{name}" deleted successfully.')
        return redirect('settings_linacs')
    
    return render(request, 'QAID_Manager/linac_confirm_delete.html', {
        'linac': linac
    })


# Physics Parameters Views
@login_required
def settings_physics_parameters(request):
    """Display list of all physics parameters (view-only for non-admin)"""
    deduplicate_film_analysis_parameters()
    physics_params = PhysicsParameters.objects.all().order_by('id')
    is_admin = request.user.is_staff
    
    return render(request, 'QAID_Manager/settings_physics_parameters.html', {
        'physics_params': physics_params,
        'is_admin': is_admin,
    })


@login_required
def physics_parameter_detail(request, pk):
    """View physics parameter details"""
    param = get_object_or_404(PhysicsParameters, pk=pk)
    film_analysis_threshold_percent = None
    film_analysis_band_width_mm = None
    if param.parameter_type == 'film_analysis' and isinstance(param.parameter_values, dict):
        threshold = param.parameter_values.get('field_size_detection_threshold')
        if threshold is not None:
            film_analysis_threshold_percent = round(float(threshold) * 100, 1)
        band_width = param.parameter_values.get('field_size_band_width_mm')
        if band_width is not None:
            film_analysis_band_width_mm = round(float(band_width), 1)

    # Extract table_data from parameter_values for easier template access
    table_data = None
    filename = None
    import_date = None
    
    if param.parameter_type != 'film_analysis' and param.parameter_values:
        # Handle dict format (wrapped or simple)
        if isinstance(param.parameter_values, dict):
            # Check for wrapped format with table_data
            if 'table_data' in param.parameter_values:
                table_data = param.parameter_values.get('table_data')
                filename = param.parameter_values.get('filename')
                import_date = param.parameter_values.get('import_date')
                # Ensure table_data is a list
                if table_data is not None:
                    if not isinstance(table_data, list):
                        table_data = [table_data] if table_data else []
            else:
                # Simple dict format - convert to list
                table_data = [param.parameter_values]
        # Handle list format
        elif isinstance(param.parameter_values, list):
            # Check if it's a list containing a dict with table_data (old format)
            if len(param.parameter_values) > 0 and isinstance(param.parameter_values[0], dict):
                first_item = param.parameter_values[0]
                if 'table_data' in first_item:
                    # Extract from wrapped format in list
                    table_data = first_item.get('table_data')
                    filename = first_item.get('filename')
                    import_date = first_item.get('import_date')
                    # Ensure table_data is a list
                    if table_data is not None and not isinstance(table_data, list):
                        table_data = [table_data] if table_data else []
                else:
                    # List of dicts (direct format)
                    table_data = param.parameter_values
            else:
                # Direct list format
                table_data = param.parameter_values
    
    # Final validation: ensure table_data is a list if it exists
    if table_data is not None and not isinstance(table_data, list):
        table_data = [table_data] if table_data else None
    
    # Only pass table_data if it's a non-empty list
    if not table_data or (isinstance(table_data, list) and len(table_data) == 0):
        table_data = None
    
    return render(request, 'QAID_Manager/physics_parameter_detail.html', {
        'param': param,
        'table_data': table_data,
        'filename': filename,
        'import_date': import_date,
        'film_analysis_threshold_percent': film_analysis_threshold_percent,
        'film_analysis_band_width_mm': film_analysis_band_width_mm,
    })


@staff_member_required
def physics_parameter_create(request):
    """Create a new physics parameter (admin only)"""
    import json
    if request.method == 'POST':
        form = PhysicsParametersForm(request.POST)
        if form.is_valid():
            param = form.save()
            messages.success(request, f'Physics Parameter "{param.name}" created successfully.')
            return redirect('settings_physics_parameters')
    else:
        form = PhysicsParametersForm()
    
    return render(request, 'QAID_Manager/physics_parameter_form.html', {
        'form': form,
        'title': 'Add New Physics Parameter',
        'parameter_values_json': json.dumps({'table_data': []})
    })


@staff_member_required
def physics_parameter_edit(request, pk):
    """Edit an existing physics parameter (admin only)"""
    import json
    param = get_object_or_404(PhysicsParameters, pk=pk)
    
    if request.method == 'POST':
        form = PhysicsParametersForm(request.POST, instance=param)
        if form.is_valid():
            param = form.save()
            messages.success(request, f'Physics Parameter "{param.name}" updated successfully.')
            return redirect('settings_physics_parameters')
    else:
        form = PhysicsParametersForm(instance=param)
    
    # Prepare parameter_values as JSON for JavaScript
    param_values_json = '{}'
    if param.parameter_values:
        if isinstance(param.parameter_values, dict) and 'table_data' in param.parameter_values:
            param_values_json = json.dumps(param.parameter_values)
        elif isinstance(param.parameter_values, list):
            param_values_json = json.dumps(param.parameter_values)
        elif isinstance(param.parameter_values, dict):
            param_values_json = json.dumps({'table_data': [param.parameter_values]})
    
    return render(request, 'QAID_Manager/physics_parameter_form.html', {
        'form': form,
        'param': param,
        'title': f'Edit Physics Parameter: {param.name}',
        'parameter_values_json': param_values_json
    })


@staff_member_required
def physics_parameter_delete(request, pk):
    """Delete a physics parameter (admin only)"""
    param = get_object_or_404(PhysicsParameters, pk=pk)
    
    if request.method == 'POST':
        name = param.name
        param.delete()
        messages.success(request, f'Physics Parameter "{name}" deleted successfully.')
        return redirect('settings_physics_parameters')
    
    return render(request, 'QAID_Manager/physics_parameter_confirm_delete.html', {
        'param': param
    })


# QA Test Views
@login_required
def settings_qa_tests(request):
    """Display list of all QA tests and user-defined test types."""
    qa_tests = QATest.objects.all().order_by('test_type', 'order_index')
    custom_types = CustomTestType.objects.prefetch_related('tests').all()
    custom_tests = CustomTest.objects.select_related('test_type').all().order_by('test_type__display_order', 'order_index')
    is_admin = request.user.is_staff
    
    return render(request, 'QAID_Manager/settings_qa_tests.html', {
        'qa_tests': qa_tests,
        'custom_types': custom_types,
        'custom_tests': custom_tests,
        'is_admin': is_admin
    })


@login_required
def qa_test_detail(request, pk):
    """View QA test details"""
    test = get_object_or_404(QATest, pk=pk)
    test_number = test.order_index if 1 <= test.order_index <= 20 else None

    placeholder_rows = []
    if test_number is not None:
        nn = f"{test_number:02d}"
        placeholder_rows = [
            {
                'field': 'Test name',
                'description': 'Name of the test',
                'value': test.name,
                'placeholder': f"{{{{TEST_{nn}_NAME}}}}",
            },
            {
                'field': 'Test value',
                'description': 'Result value of the test',
                'value': 'e.g. 1.0',
                'placeholder': f"{{{{TEST_{nn}_VALUE}}}}",
            },
            {
                'field': 'Test tolerance',
                'description': 'Tolerance used for validation',
                'value': f"±{test.tolerance_value} {test.tolerance_unit}",
                'placeholder': f"{{{{TEST_{nn}_TOLERANCE}}}}",
            },
            {
                'field': 'Test status',
                'description': 'Pass/fail result text',
                'value': 'Đạt / Không đạt',
                'placeholder': f"{{{{TEST_{nn}_STATUS}}}}",
            },
        ]
    return render(request, 'QAID_Manager/qa_test_detail.html', {
        'test': test,
        'placeholder_rows': placeholder_rows,
        'test_number': test_number,
    })


@staff_member_required
def qa_test_create(request):
    """Create a new QA test (admin only). Routes to CustomTest when a user-defined type is selected."""
    if request.method == 'POST':
        selected_type = request.POST.get('test_type', '')
        if selected_type.startswith('custom_'):
            slug = selected_type[len('custom_'):]
            ct = CustomTestType.objects.filter(slug=slug).first()
            if ct:
                from ..forms import CustomTestForm as CTForm
                ct_form = CTForm(request.POST)
                if ct_form.is_valid():
                    test = ct_form.save(commit=False)
                    test.test_type = ct
                    if not test.order_index:
                        test.order_index = (ct.tests.aggregate(m=Max('order_index'))['m'] or 0) + 1
                    test.save()
                    messages.success(request, f'Test "{test.name}" added to {ct.name}.')
                    return redirect('settings_qa_tests')
                else:
                    form = QATestForm(request.POST)
                    return render(request, 'QAID_Manager/qa_test_form.html', {
                        'form': form,
                        'title': 'Add New QA Test'
                    })
            else:
                messages.error(request, 'Selected test type not found.')
                form = QATestForm(request.POST)
        else:
            form = QATestForm(request.POST)
            if form.is_valid():
                test = form.save()
                messages.success(request, f'QA Test "{test.name}" created successfully.')
                return redirect('settings_qa_tests')
    else:
        form = QATestForm()
    
    return render(request, 'QAID_Manager/qa_test_form.html', {
        'form': form,
        'title': 'Add New QA Test'
    })


@staff_member_required
def qa_test_edit(request, pk):
    """Edit an existing QA test (admin only)"""
    test = get_object_or_404(QATest, pk=pk)
    
    if request.method == 'POST':
        form = QATestForm(request.POST, instance=test)
        if form.is_valid():
            test = form.save()
            messages.success(request, f'QA Test "{test.name}" updated successfully.')
            return redirect('settings_qa_tests')
    else:
        form = QATestForm(instance=test)
    
    return render(request, 'QAID_Manager/qa_test_form.html', {
        'form': form,
        'test': test,
        'title': f'Edit QA Test: {test.name}'
    })


@staff_member_required
def qa_test_delete(request, pk):
    """Delete a QA test (admin only)"""
    test = get_object_or_404(QATest, pk=pk)
    
    if request.method == 'POST':
        name = test.name
        test.delete()
        messages.success(request, f'QA Test "{name}" deleted successfully.')
        return redirect('settings_qa_tests')
    
    return render(request, 'QAID_Manager/qa_test_confirm_delete.html', {
        'test': test
    })


# Vietnamese Holiday Views
@login_required
def settings_vietnamese_holidays(request):
    """Display list of all Vietnamese holidays (view-only for non-admin)"""
    holidays = VietnameseHoliday.objects.all().order_by('date')
    is_admin = request.user.is_staff
    
    return render(request, 'QAID_Manager/settings_vietnamese_holidays.html', {
        'holidays': holidays,
        'is_admin': is_admin
    })


@login_required
def vietnamese_holiday_detail(request, pk):
    """View Vietnamese holiday details"""
    holiday = get_object_or_404(VietnameseHoliday, pk=pk)
    return render(request, 'QAID_Manager/vietnamese_holiday_detail.html', {
        'holiday': holiday
    })


@staff_member_required
def vietnamese_holiday_create(request):
    """Create a new Vietnamese holiday (admin only)"""
    if request.method == 'POST':
        form = VietnameseHolidayForm(request.POST)
        if form.is_valid():
            holiday = form.save()
            messages.success(request, f'Holiday "{holiday.name}" created successfully.')
            return redirect('settings_vietnamese_holidays')
    else:
        form = VietnameseHolidayForm()
    
    return render(request, 'QAID_Manager/vietnamese_holiday_form.html', {
        'form': form,
        'title': 'Add New Holiday'
    })


@staff_member_required
def vietnamese_holiday_edit(request, pk):
    """Edit an existing Vietnamese holiday (admin only)"""
    holiday = get_object_or_404(VietnameseHoliday, pk=pk)
    
    if request.method == 'POST':
        form = VietnameseHolidayForm(request.POST, instance=holiday)
        if form.is_valid():
            holiday = form.save()
            messages.success(request, f'Holiday "{holiday.name}" updated successfully.')
            return redirect('settings_vietnamese_holidays')
    else:
        form = VietnameseHolidayForm(instance=holiday)
    
    return render(request, 'QAID_Manager/vietnamese_holiday_form.html', {
        'form': form,
        'holiday': holiday,
        'title': f'Edit Holiday: {holiday.name}'
    })


@staff_member_required
def vietnamese_holiday_delete(request, pk):
    """Delete a Vietnamese holiday (admin only)"""
    holiday = get_object_or_404(VietnameseHoliday, pk=pk)
    
    if request.method == 'POST':
        name = holiday.name
        holiday.delete()
        messages.success(request, f'Holiday "{name}" deleted successfully.')
        return redirect('settings_vietnamese_holidays')
    
    return render(request, 'QAID_Manager/vietnamese_holiday_confirm_delete.html', {
        'holiday': holiday
    })


# Organization Settings Views
@login_required
@staff_member_required
def settings_organization(request):
    """Edit organization settings (logo and organization name) - admin only"""
    settings = OrganizationSettings.get_settings()
    
    if request.method == 'POST':
        form = OrganizationSettingsForm(request.POST, request.FILES, instance=settings)
        if form.is_valid():
            # Keep a stable filename for templates so each upload overwrites
            # the previous one instead of creating qa_report_template_1, _2, ...
            template_uploads = {
                'report_template': 'organization/reports/qa_report_template.docx',
                'service_report_template': 'organization/reports/service_report_template.docx',
            }
            for field_name, fixed_path in template_uploads.items():
                uploaded = request.FILES.get(field_name)
                if not uploaded:
                    continue
                # Remove existing fixed target path first so FileSystemStorage
                # does not append suffixes for "available" names.
                if settings.__class__._meta.get_field(field_name).storage.exists(fixed_path):
                    settings.__class__._meta.get_field(field_name).storage.delete(fixed_path)
                uploaded.name = os.path.basename(fixed_path)

            form.save()
            messages.success(request, 'Organization settings updated successfully.')
            return redirect('settings_organization')
    else:
        form = OrganizationSettingsForm(instance=settings)
    
    return render(request, 'QAID_Manager/settings_organization.html', {
        'form': form,
        'settings': settings
    })


@login_required
@staff_member_required
def download_report_template(request, template_type):
    """
    Download default QA/service report template.
    Falls back to generated sample template when bundled file is missing.
    """
    template_map = {
        'qa': 'qa_report_template.docx',
        'service': 'service_report_template.docx',
    }
    if template_type not in template_map:
        raise Http404("Unknown template type.")

    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    template_path = os.path.join(base_dir, 'QAID_Manager', 'templates', 'reports', template_map[template_type])

    if os.path.exists(template_path):
        return FileResponse(
            open(template_path, 'rb'),
            as_attachment=True,
            filename=template_map[template_type],
        )

    from docx import Document
    doc = Document()
    if template_type == 'qa':
        doc.add_heading('QA Report Template', level=1)
        doc.add_paragraph('LINAC: {{LINAC_NAME}}')
        doc.add_paragraph('Date: {{DATE_PERFORMED}}')
        doc.add_paragraph('Performed by: {{PERFORMED_BY}}')
        doc.add_paragraph('Test 01 Value: {{TEST_01_VALUE}}')
        doc.add_paragraph('Test 01 Tolerance: {{TEST_01_TOLERANCE}}')
        filename = 'qa_report_template_sample.docx'
    else:
        doc.add_heading('Service Report Template', level=1)
        doc.add_paragraph('Ngày in: {{REPORT_PRINTING_DATE}} / {{REPORT_PRINTING_MONTH}} / {{REPORT_PRINTING_YEAR}}')
        doc.add_paragraph('Chèn bảng dịch vụ tại đây:')
        doc.add_paragraph('{{SERVICE_TABLE}}')
        filename = 'service_report_template_sample.docx'

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename=filename)


# User Management Views
@login_required
@staff_member_required
def settings_users(request):
    """Display list of all users (admin only)"""
    users = User.objects.all().order_by('username')
    
    return render(request, 'QAID_Manager/settings_users.html', {
        'users': users
    })


@login_required
@staff_member_required
def user_detail(request, pk):
    """View user details (admin only)"""
    user = get_object_or_404(User, pk=pk)
    # Get all activities for this user, ordered by most recent first
    activities = UserActivity.objects.filter(user=user).order_by('-created_at')[:100]  # Limit to last 100 activities
    
    return render(request, 'QAID_Manager/user_detail.html', {
        'user': user,
        'activities': activities
    })


@login_required
@staff_member_required
def user_create(request):
    """Create a new user (admin only)"""
    if request.method == 'POST':
        form = UserForm(request.POST)
        if form.is_valid():
            user = form.save()
            messages.success(request, f'User "{user.username}" created successfully.')
            return redirect('settings_users')
    else:
        form = UserForm()
    
    return render(request, 'QAID_Manager/user_form.html', {
        'form': form,
        'title': 'Add New User'
    })


@login_required
@staff_member_required
def user_edit(request, pk):
    """Edit an existing user (admin only)"""
    user = get_object_or_404(User, pk=pk)
    
    if request.method == 'POST':
        form = UserForm(request.POST, instance=user)
        if form.is_valid():
            user = form.save()
            messages.success(request, f'User "{user.username}" updated successfully.')
            return redirect('settings_users')
    else:
        form = UserForm(instance=user)
    
    return render(request, 'QAID_Manager/user_form.html', {
        'form': form,
        'user': user,
        'title': f'Edit User: {user.username}'
    })


@login_required
@staff_member_required
def user_delete(request, pk):
    """Delete a user (admin only)"""
    user = get_object_or_404(User, pk=pk)
    
    # Prevent deleting yourself
    if user == request.user:
        messages.error(request, 'You cannot delete your own account.')
        return redirect('settings_users')
    
    if request.method == 'POST':
        username = user.username
        user.delete()
        messages.success(request, f'User "{username}" deleted successfully.')
        return redirect('settings_users')
    
    return render(request, 'QAID_Manager/user_confirm_delete.html', {
        'user': user
    })


# Device/Equipment Views
@login_required
def settings_devices(request):
    """Display list of all devices (view-only for non-admin)"""
    devices = Device.objects.all().order_by('category', 'name')
    is_admin = request.user.is_staff
    
    return render(request, 'QAID_Manager/settings_devices.html', {
        'devices': devices,
        'is_admin': is_admin
    })


@login_required
def device_detail(request, pk):
    """View device details"""
    device = get_object_or_404(Device, pk=pk)
    documents = device.documents.all()
    is_admin = request.user.is_staff
    
    # Handle file upload
    if request.method == 'POST' and is_admin:
        form = DeviceDocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.device = device
            document.uploaded_by = request.user
            
            # Determine file type from extension
            filename = document.file.name
            ext = filename.split('.')[-1].lower() if '.' in filename else ''
            file_type_map = {
                'pdf': 'pdf',
                'doc': 'doc',
                'docx': 'docx',
                'png': 'png',
                'jpg': 'jpg',
                'jpeg': 'jpeg',
                'tiff': 'tiff',
                'tif': 'tif',
            }
            document.file_type = file_type_map.get(ext, 'other')
            
            # Set file_name if not provided
            if not document.file_name:
                document.file_name = os.path.basename(filename)
            
            document.save()
            messages.success(request, f'File "{document.file_name}" uploaded successfully.')
            return redirect('device_detail', pk=pk)
        else:
            messages.error(request, 'Error uploading file. Please check the form.')
    else:
        form = DeviceDocumentForm()
    
    return render(request, 'QAID_Manager/device_detail.html', {
        'device': device,
        'documents': documents,
        'form': form,
        'is_admin': is_admin
    })


@login_required
@staff_member_required
@require_POST
def device_document_delete(request, pk, doc_pk):
    """Delete a device document"""
    document = get_object_or_404(DeviceDocument, pk=doc_pk, device_id=pk)
    file_name = document.file_name
    document.delete()
    messages.success(request, f'File "{file_name}" deleted successfully.')
    return redirect('device_detail', pk=pk)


@staff_member_required
def device_create(request):
    """Create a new device (admin only)"""
    if request.method == 'POST':
        form = DeviceForm(request.POST)
        if form.is_valid():
            device = form.save()
            messages.success(request, f'Device "{device.name}" created successfully.')
            return redirect('settings_devices')
    else:
        form = DeviceForm()
    
    return render(request, 'QAID_Manager/device_form.html', {
        'form': form,
        'title': 'Add New Device'
    })


@staff_member_required
def device_edit(request, pk):
    """Edit an existing device (admin only)"""
    device = get_object_or_404(Device, pk=pk)
    
    if request.method == 'POST':
        form = DeviceForm(request.POST, instance=device)
        if form.is_valid():
            device = form.save()
            messages.success(request, f'Device "{device.name}" updated successfully.')
            return redirect('settings_devices')
    else:
        form = DeviceForm(instance=device)
    
    return render(request, 'QAID_Manager/device_form.html', {
        'form': form,
        'device': device,
        'title': f'Edit Device: {device.name}'
    })


@staff_member_required
def device_delete(request, pk):
    """Delete a device (admin only)"""
    device = get_object_or_404(Device, pk=pk)
    
    if request.method == 'POST':
        name = device.name
        device.delete()
        messages.success(request, f'Device "{name}" deleted successfully.')
        return redirect('settings_devices')
    
    return render(request, 'QAID_Manager/device_confirm_delete.html', {
        'device': device
    })


# ============================================================================
# Custom Test Types Management
# ============================================================================

@login_required
def settings_custom_test_types(request):
    """Redirect to QA Tests page which now includes test type management."""
    return redirect('settings_qa_tests')


@staff_member_required
def custom_test_type_create(request):
    """Create a new custom test type."""
    if request.method == 'POST':
        form = CustomTestTypeForm(request.POST)
        if form.is_valid():
            from django.utils.text import slugify
            ct = form.save(commit=False)
            ct.slug = slugify(ct.name) or 'custom'
            base_slug = ct.slug
            n = 1
            while CustomTestType.objects.filter(slug=ct.slug).exists():
                ct.slug = f'{base_slug}-{n}'
                n += 1
            if not ct.display_order:
                max_order = CustomTestType.objects.aggregate(m=Max('display_order'))['m'] or 0
                ct.display_order = max_order + 1
            ct.save()
            messages.success(request, f'Test type "{ct.name}" created.')
            return redirect('settings_qa_tests')
    else:
        next_order = (CustomTestType.objects.aggregate(m=Max('display_order'))['m'] or 0) + 1
        form = CustomTestTypeForm(initial={'display_order': next_order})
    return render(request, 'QAID_Manager/custom_test_type_form.html', {
        'form': form,
        'title': 'Add New Test Type',
    })


@staff_member_required
def custom_test_type_edit(request, pk):
    """Edit a custom test type."""
    ct = get_object_or_404(CustomTestType, pk=pk)
    if request.method == 'POST':
        form = CustomTestTypeForm(request.POST, instance=ct)
        if form.is_valid():
            form.save()
            messages.success(request, f'Test type "{ct.name}" updated.')
            return redirect('settings_qa_tests')
    else:
        form = CustomTestTypeForm(instance=ct)
    return render(request, 'QAID_Manager/custom_test_type_form.html', {
        'form': form,
        'title': f'Edit: {ct.name}',
        'custom_type': ct,
    })


@staff_member_required
def custom_test_type_delete(request, pk):
    """Delete a custom test type."""
    ct = get_object_or_404(CustomTestType, pk=pk)
    if request.method == 'POST':
        name = ct.name
        ct.delete()
        messages.success(request, f'Test type "{name}" deleted.')
        return redirect('settings_qa_tests')
    return render(request, 'QAID_Manager/custom_test_type_confirm_delete.html', {
        'custom_type': ct,
    })


@staff_member_required
def custom_test_create(request, type_pk):
    """Create a new test within a custom test type."""
    ct = get_object_or_404(CustomTestType, pk=type_pk)
    if request.method == 'POST':
        form = CustomTestForm(request.POST)
        if form.is_valid():
            test = form.save(commit=False)
            test.test_type = ct
            test.save()
            messages.success(request, f'Test "{test.name}" added to {ct.name}.')
            return redirect('settings_qa_tests')
    else:
        next_order = (ct.tests.count() or 0) + 1
        form = CustomTestForm(initial={'order_index': next_order})
    return render(request, 'QAID_Manager/custom_test_form.html', {
        'form': form,
        'custom_type': ct,
        'title': f'Add Test to {ct.name}',
    })


@staff_member_required
def custom_test_edit(request, pk):
    """Edit a custom test."""
    test = get_object_or_404(CustomTest, pk=pk)
    ct = test.test_type
    if request.method == 'POST':
        form = CustomTestForm(request.POST, instance=test)
        if form.is_valid():
            form.save()
            messages.success(request, f'Test "{test.name}" updated.')
            return redirect('settings_qa_tests')
    else:
        form = CustomTestForm(instance=test)
    return render(request, 'QAID_Manager/custom_test_form.html', {
        'form': form,
        'custom_type': ct,
        'title': f'Edit: {test.name}',
        'custom_test': test,
    })


@staff_member_required
def custom_test_delete(request, pk):
    """Delete a custom test."""
    test = get_object_or_404(CustomTest, pk=pk)
    if request.method == 'POST':
        name = test.name
        test.delete()
        messages.success(request, f'Test "{name}" deleted.')
        return redirect('settings_qa_tests')
    return render(request, 'QAID_Manager/custom_test_confirm_delete.html', {
        'custom_test': test,
    })
