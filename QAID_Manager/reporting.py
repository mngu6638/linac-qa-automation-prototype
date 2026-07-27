"""
QA Report Generation Module.

This module provides functionality to generate DOCX and PDF reports
for QA records, including test results, dose calculations, film analyses,
and test notes.
"""
import os
import tempfile
import logging
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.conf import settings
from django.http import HttpResponse
from .models import QARecord, DoseCalculation, QAFilmAnalysis, QATestNote
from .export_converter import convert_docx_to_pdf

logger = logging.getLogger(__name__)

# ============================================================================
# QA Report Generator Class
# ============================================================================

class QAReportGenerator:
    """
    Generate DOCX and PDF reports for QA records.
    
    Supports custom templates from organization settings or default templates.
    Replaces placeholders in templates with actual QA data and generates
    comprehensive reports including test results, dose calculations, and analyses.
    """
    
    def __init__(self):
        self.template_dir = os.path.join(settings.BASE_DIR, 'QAID_Manager', 'templates', 'reports')
        os.makedirs(self.template_dir, exist_ok=True)
    
    def generate_qa_report(self, qa_record, report_type='standard'):
        """Generate a DOCX report for a QA record"""
        # Try to use organization settings template first
        from .models import OrganizationSettings
        org_settings = OrganizationSettings.get_settings()
        template_path = None
        
        if org_settings.report_template and org_settings.report_template.name:
            # Use template from organization settings
            template_path = org_settings.report_template.path
        else:
            # Fall back to file system template
            template_path = os.path.join(self.template_dir, 'qa_report_template.docx')
        
        if template_path and os.path.exists(template_path):
            # Use the custom template as starting point
            doc = Document(template_path)
            # Replace placeholders in the document
            self._replace_placeholders(doc, qa_record)
        else:
            # Create new document if no template exists
            doc = Document()
            
            # Set up document properties
            doc.core_properties.title = f"QA Report - {qa_record.linac.name} - {qa_record.date_performed}"
            doc.core_properties.author = qa_record.performed_by.username if qa_record.performed_by else "System"
            doc.core_properties.created = datetime.now()
            
            # Add header
            self._add_header(doc, qa_record)
            
            # Add basic information
            self._add_basic_info(doc, qa_record)
            
            # Add test results
            self._add_test_results(doc, qa_record)
            
            # Add dose calculations if available
            dose_calculations = qa_record.dose_calculations.all()
            if dose_calculations.exists():
                self._add_dose_calculations(doc, dose_calculations)
            
            # Add film analyses if available
            film_analyses = qa_record.film_analyses.all()
            if film_analyses.exists():
                self._add_film_analyses(doc, film_analyses)
            
            # Add test notes if available
            test_notes = qa_record.test_notes.all()
            if test_notes.exists():
                self._add_test_notes(doc, test_notes)
            
            # Add summary and conclusions
            self._add_summary(doc, qa_record)
        
        return doc
    
    def _replace_placeholders(self, doc, qa_record):
        """Replace placeholders in the template with actual data"""
        # Get user's full name instead of username
        user_full_name = "N/A"
        if qa_record.performed_by:
            if hasattr(qa_record.performed_by, 'first_name') and hasattr(qa_record.performed_by, 'last_name'):
                first_name = qa_record.performed_by.first_name or ""
                last_name = qa_record.performed_by.last_name or ""
                if first_name or last_name:
                    user_full_name = f"{last_name} {first_name}".strip()
                else:
                    user_full_name = qa_record.performed_by.username
            else:
                user_full_name = qa_record.performed_by.username
        
        # Convert UTC time to local time (using Django's current timezone setting)
        from django.utils import timezone
        
        # Use Django's current timezone (same as what Django admin uses)
        if timezone.is_aware(qa_record.created_at):
            local_time = qa_record.created_at.astimezone(timezone.get_current_timezone())
        else:
            # If naive datetime, assume it's already in local time
            local_time = qa_record.created_at
        
        # Option: Use a fixed time for the report (e.g., 2:00 PM)
        # Uncomment the next line if you want to always show 2:00 PM
        # local_time = local_time.replace(hour=14, minute=0, second=0, microsecond=0)
        
        # Define the replacements
        replacements = {
            '{{LINAC_NAME}}': qa_record.linac.name,
            '{{DATE_PERFORMED}}': qa_record.date_performed.strftime('%Y-%m-%d'),
            '{{PERFORMED_BY}}': user_full_name,
            '{{STATUS}}': qa_record.status.name if qa_record.status else 'N/A',
            # New specific placeholders
            '{{QA_PERFORMED_DATE_DD}}': qa_record.date_performed.strftime('%d'),
            '{{QA_PERFORMED_DATE_MM}}': qa_record.date_performed.strftime('%m'),
            '{{QA_PERFORMED_DATE_YYYY}}': qa_record.date_performed.strftime('%Y'),
            '{{QA_PERFORMED_DATE_HHMM}}': local_time.strftime('%H:%M'),  # Use local time
            '{{QA_PERFORMED_BY_USER}}': user_full_name,
            '{{LINAC_NAME}}': qa_record.linac.name,
            '{{LINAC_SERIES_NUMBER}}': qa_record.linac.series_number if hasattr(qa_record.linac, 'series_number') else 'N/A',
        }
        
        # Add individual test placeholders
        test_names = {
            'test_01': 'Kích thước trường ánh sáng (đối xứng và bất đối xứng)',
            'test_02': 'Góc quay bộ chuẩn trực (collimator)',
            'test_03': 'Góc quay bàn điều trị',
            'test_04': 'Độ chính xác trong chuyển động của Bàn điều trị',
            'test_05': 'Góc quay thân máy (gantry)',
            'test_06': 'Độ chính xác của chùm laser tại điểm đồng tâm',
            'test_07': 'Độ chính xác của ODI',
            'test_08': 'Tâm dây chữ thập',
            'test_09': 'Đồng tâm quay collimator',
            'test_10': 'Đồng tâm quay bàn điều trị',
            'test_11': 'Đồng tâm quay gantry',
            'test_12': 'Field Size Analysis',
            'test_13': 'Collimator Isocenter Circle Diameter',
            'test_14': 'Gantry Isocenter Circle Diameter',
            'test_15': 'Liều tuyệt đối',
            'test_16': 'Sự ổn định của năng lượng (D10)',
            'test_17': 'Tính đối xứng so với giá trị tại thời điểm commissioning',
            'test_18': 'Tính phẳng so với giá trị tại thời điểm commissioning',
            'test_19': 'Tính tuyến tính của liều và số MU',
            'test_20': 'Hệ số liều lối ra theo kích thước trường chiếu',
        }
        
        # Add individual test value placeholders
        for i in range(1, 21):
            field_name = f'test_{i:02d}'
            value = getattr(qa_record, field_name)
            test_name = test_names.get(field_name, f'Test {i}')
            
            # Add test value placeholder (without curly braces)
            replacements[f'{{TEST_{i:02d}_VALUE}}'] = f"{value:.2f}" if value is not None else "N/A"
            # Add test name placeholder
            replacements[f'{{TEST_{i:02d}_NAME}}'] = test_name
            # Add test status placeholder (Vietnamese and without curly braces)
            tolerance = 2.0 if i in [11, 15, 20] else 1.0 if i in [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 12, 13, 14, 16, 19] else 3.0 if i in [17, 18] else 0
            if value is None:
                status = "Không có dữ liệu"
            elif abs(value) <= tolerance:
                status = "Đạt"
            else:
                status = "Không đạt"
            replacements[f'{{TEST_{i:02d}_STATUS}}'] = status
            # Add test tolerance placeholder
            replacements[f'{{TEST_{i:02d}_TOLERANCE}}'] = f"±{tolerance}"
        
        # Replace placeholders in paragraphs while preserving formatting
        def replace_in_paragraph(paragraph):
            """Replace placeholders in a paragraph while preserving formatting"""
            # Replace placeholders in each run individually to preserve formatting
            for run in paragraph.runs:
                original_text = run.text
                modified_text = original_text
                
                # Replace all placeholders (support both {{}} and {} formats)
                for placeholder, replacement in replacements.items():
                    # Replace double braces format
                    if placeholder in modified_text:
                        modified_text = modified_text.replace(placeholder, replacement)
                    # Also support single braces format
                    single_brace_placeholder = placeholder.replace('{{', '{').replace('}}', '}')
                    if single_brace_placeholder in modified_text:
                        modified_text = modified_text.replace(single_brace_placeholder, replacement)
                
                # Only update if text actually changed
                if modified_text != original_text:
                    run.text = modified_text
        
        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        
        # Add dynamic content sections
        self._add_dynamic_sections(doc, qa_record)
    
    def _add_dynamic_sections(self, doc, qa_record):
        """Add dynamic content sections to the template"""
        # Find where to add dynamic content
        for paragraph in doc.paragraphs:
            if '{{TEST_RESULTS}}' in paragraph.text:
                # Clear the placeholder
                paragraph.clear()
                # Add test results
                self._add_test_results_to_paragraph(paragraph, qa_record)
            
            elif '{{DOSE_CALCULATIONS}}' in paragraph.text:
                paragraph.clear()
                dose_calculations = qa_record.dose_calculations.all()
                if dose_calculations.exists():
                    self._add_dose_calculations_to_paragraph(paragraph, dose_calculations)
            
            elif '{{FILM_ANALYSES}}' in paragraph.text:
                paragraph.clear()
                film_analyses = qa_record.film_analyses.all()
                if film_analyses.exists():
                    self._add_film_analyses_to_paragraph(paragraph, film_analyses)
            
            elif '{{TEST_NOTES}}' in paragraph.text:
                paragraph.clear()
                test_notes = qa_record.test_notes.all()
                if test_notes.exists():
                    self._add_test_notes_to_paragraph(paragraph, test_notes)
            
            elif '{{SUMMARY}}' in paragraph.text:
                paragraph.clear()
                self._add_summary_to_paragraph(paragraph, qa_record)
    
    def _add_test_results_to_paragraph(self, paragraph, qa_record):
        """Add test results to a paragraph"""
        paragraph.add_run("Kết quả kiểm tra:\n").bold = True
        
        # Get all test results
        test_results = []
        for field in qa_record._meta.fields:
            if field.name.startswith('test_') and field.name != 'test_notes':
                value = getattr(qa_record, field.name)
                if value is not None:
                    test_results.append((field.verbose_name, value))
        
        for test_name, value in test_results:
            paragraph.add_run(f"• {test_name}: {value:.2f}\n")
    
    def _add_dose_calculations_to_paragraph(self, paragraph, dose_calculations):
        """Add dose calculations to a paragraph"""
        paragraph.add_run("Tính toán liều:\n").bold = True
        
        for calc in dose_calculations:
            paragraph.add_run(f"• Năng lượng: {calc.energy}\n")
            paragraph.add_run(f"• Liều tuyệt đối: {calc.absolute_dose_deviation:.3f}%\n")
            paragraph.add_run(f"• Năng lượng chùm tia: {calc.beam_energy_d10:.3f}%\n")
    
    def _add_film_analyses_to_paragraph(self, paragraph, film_analyses):
        """Add film analyses to a paragraph"""
        paragraph.add_run("Phân tích phim:\n").bold = True
        
        for analysis in film_analyses:
            paragraph.add_run(f"• {analysis.film_type}: {analysis.result}\n")
    
    def _add_test_notes_to_paragraph(self, paragraph, test_notes):
        """Add test notes to a paragraph"""
        paragraph.add_run("Ghi chú kiểm tra:\n").bold = True
        
        for note in test_notes:
            paragraph.add_run(f"• {note.note}\n")
    
    def _add_summary_to_paragraph(self, paragraph, qa_record):
        """Add summary to a paragraph"""
        paragraph.add_run("Tóm tắt:\n").bold = True
        
        # Count pass/fail tests
        total_tests = 0
        passed_tests = 0
        
        for field in qa_record._meta.fields:
            if field.name.startswith('test_') and field.name != 'test_notes':
                value = getattr(qa_record, field.name)
                if value is not None:
                    total_tests += 1
                    # Simple pass/fail logic (you can customize this)
                    if abs(value) <= 2.0:  # Assuming 2% tolerance
                        passed_tests += 1
        
        paragraph.add_run(f"Tổng số kiểm tra: {total_tests}\n")
        paragraph.add_run(f"Đạt: {passed_tests}\n")
        paragraph.add_run(f"Không đạt: {total_tests - passed_tests}\n")
        
        if passed_tests == total_tests:
            paragraph.add_run("Tình trạng tổng thể: Đạt\n").bold = True
        else:
            paragraph.add_run("Tình trạng tổng thể: Không đạt\n").bold = True
    
    def _add_header(self, doc, qa_record):
        """Add report header"""
        # Title
        title = doc.add_heading('QA Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"LINAC: {qa_record.linac.name}").bold = True
        subtitle.add_run(f" | Date: {qa_record.date_performed.strftime('%Y-%m-%d')}").bold = True
        
        doc.add_paragraph()  # Spacing
    
    def _add_basic_info(self, doc, qa_record):
        """Add basic QA record information"""
        doc.add_heading('Basic Information', level=1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Value'
        
        # Get user's full name
        user_full_name = "N/A"
        if qa_record.performed_by:
            if hasattr(qa_record.performed_by, 'first_name') and hasattr(qa_record.performed_by, 'last_name'):
                first_name = qa_record.performed_by.first_name or ""
                last_name = qa_record.performed_by.last_name or ""
                if first_name or last_name:
                    user_full_name = f"{last_name} {first_name}".strip()
                else:
                    user_full_name = qa_record.performed_by.username
            else:
                user_full_name = qa_record.performed_by.username
        
        # Data rows
        data = [
            ('LINAC', qa_record.linac.name),
            ('Date Performed', qa_record.date_performed.strftime('%Y-%m-%d')),
            ('Performed By', user_full_name),
            ('Status', qa_record.status.name if qa_record.status else 'N/A'),
            ('Created', qa_record.created_at.strftime('%Y-%m-%d %H:%M')),
        ]
        
        for field, value in data:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # Spacing
    
    def _add_test_results(self, doc, qa_record):
        """Add test results table"""
        doc.add_heading('Test Results', level=1)
        
        # Test name mapping
        test_names = {
            'test_01': 'Kích thước trường ánh sáng (đối xứng và bất đối xứng)',
            'test_02': 'Góc quay bộ chuẩn trực (collimator)',
            'test_03': 'Góc quay bàn điều trị',
            'test_04': 'Độ chính xác trong chuyển động của Bàn điều trị',
            'test_05': 'Góc quay thân máy (gantry)',
            'test_06': 'Độ chính xác của chùm laser tại điểm đồng tâm',
            'test_07': 'Độ chính xác của ODI',
            'test_08': 'Tâm dây chữ thập',
            'test_09': 'Đồng tâm quay collimator',
            'test_10': 'Đồng tâm quay bàn điều trị',
            'test_11': 'Đồng tâm quay gantry',
            'test_12': 'Field Size Analysis',
            'test_13': 'Collimator Isocenter Circle Diameter',
            'test_14': 'Gantry Isocenter Circle Diameter',
            'test_15': 'Liều tuyệt đối',
            'test_16': 'Sự ổn định của năng lượng (D10)',
            'test_17': 'Tính đối xứng so với giá trị tại thời điểm commissioning',
            'test_18': 'Tính phẳng so với giá trị tại thời điểm commissioning',
            'test_19': 'Tính tuyến tính của liều và số MU',
            'test_20': 'Hệ số liều lối ra theo kích thước trường chiếu',
        }
        
        # Tolerance mapping
        tolerances = {
            'test_01': 1.0, 'test_02': 1.0, 'test_03': 1.0, 'test_04': 1.0, 'test_05': 1.0,
            'test_06': 1.0, 'test_07': 1.0, 'test_08': 1.0, 'test_09': 1.0, 'test_10': 1.0,
            'test_11': 2.0, 'test_12': 1.0, 'test_13': 1.0, 'test_14': 1.0, 'test_15': 2.0,
            'test_16': 1.0, 'test_17': 3.0, 'test_18': 3.0, 'test_19': 1.0, 'test_20': 2.0,
        }
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Kiểm tra'
        header_cells[1].text = 'Giá trị'
        header_cells[2].text = 'Dung sai'
        header_cells[3].text = 'Tình trạng'
        
        # Data rows
        for i in range(1, 21):
            field_name = f'test_{i:02d}'
            value = getattr(qa_record, field_name)
            test_name = test_names.get(field_name, f'Test {i}')
            tolerance = tolerances.get(field_name, 0)
            
            row_cells = table.add_row().cells
            row_cells[0].text = test_name
            row_cells[1].text = f"{value:.2f}" if value is not None else "N/A"
            row_cells[2].text = f"±{tolerance}"
            
            # Status (Vietnamese)
            if value is None:
                row_cells[3].text = "Không có dữ liệu"
            elif abs(value) <= tolerance:
                row_cells[3].text = "Đạt"
            else:
                row_cells[3].text = "Không đạt"
        
        doc.add_paragraph()  # Spacing
    
    def _add_dose_calculations(self, doc, dose_calculations):
        """Add dose calculation results"""
        doc.add_heading('Dose Calculations', level=1)
        
        for calc in dose_calculations:
            # General Information
            doc.add_heading(f'Dose Calculation - {calc.created_at.strftime("%Y-%m-%d %H:%M")}', level=2)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Parameter'
            header_cells[1].text = 'Value'
            
            # General info
            general_data = [
                ('LINAC', calc.linac.name),
                ('Energy', calc.energy),
                ('Detector', f"{calc.detector.name} (SN: {calc.detector.series_number})"),
                ('Phantom', calc.phantom),
            ]
            
            for param, value in general_data:
                row_cells = table.add_row().cells
                row_cells[0].text = param
                row_cells[1].text = str(value)
            
            # Absolute dose results
            doc.add_heading('Absolute Dose Results', level=3)
            abs_table = doc.add_table(rows=1, cols=2)
            abs_table.style = 'Table Grid'
            
            abs_header = abs_table.rows[0].cells
            abs_header[0].text = 'Parameter'
            abs_header[1].text = 'Value'
            
            abs_data = [
                ('Raw Measurement', f"{calc.raw_measurement} nC"),
                ('Ktp', f"{calc.ktp_result:.4f}"),
                ('Kpol', f"{calc.kpol_result:.4f}"),
                ('Ks', f"{calc.ks_result:.4f}"),
                ('MQ', f"{calc.mq_result:.2f} nC"),
                ('Dw,Q(zref)', f"{calc.dwq_zref:.2f} cGy"),
                ('Dw,Q(zmax)', f"{calc.dwq_zmax:.2f} cGy"),
            ]
            
            for param, value in abs_data:
                row_cells = abs_table.add_row().cells
                row_cells[0].text = param
                row_cells[1].text = value
            
            # Relative dose results
            if any([calc.symmetry_crossline, calc.symmetry_inline, calc.flatness_crossline, calc.flatness_inline]):
                doc.add_heading('Relative Dose Results', level=3)
                rel_table = doc.add_table(rows=1, cols=2)
                rel_table.style = 'Table Grid'
                
                rel_header = rel_table.rows[0].cells
                rel_header[0].text = 'Parameter'
                rel_header[1].text = 'Value'
                
                rel_data = []
                if calc.symmetry_crossline:
                    rel_data.append(('Symmetry Crossline', f"{calc.symmetry_crossline:.2f}%"))
                if calc.symmetry_inline:
                    rel_data.append(('Symmetry Inline', f"{calc.symmetry_inline:.2f}%"))
                if calc.flatness_crossline:
                    rel_data.append(('Flatness Crossline', f"{calc.flatness_crossline:.2f}%"))
                if calc.flatness_inline:
                    rel_data.append(('Flatness Inline', f"{calc.flatness_inline:.2f}%"))
                if calc.output_factor:
                    rel_data.append(('Output Factor', f"{calc.output_factor:.3f}"))
                if calc.wedge_factor:
                    rel_data.append(('Wedge Factor', f"{calc.wedge_factor:.3f}"))
                if calc.beam_energy_d10:
                    rel_data.append(('Beam Energy (D10)', f"{calc.beam_energy_d10:.3f}"))
                
                for param, value in rel_data:
                    row_cells = rel_table.add_row().cells
                    row_cells[0].text = param
                    row_cells[1].text = value
            
            doc.add_paragraph()  # Spacing
    
    def _add_film_analyses(self, doc, film_analyses):
        """Add film analysis results"""
        doc.add_heading('Film Analysis Results', level=1)
        
        for analysis in film_analyses:
            doc.add_heading(f'{analysis.analysis_type.replace("_", " ").title()}', level=2)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Field'
            header_cells[1].text = 'Value'
            
            data = [
                ('Analysis Type', analysis.analysis_type),
                ('Created', analysis.created_at.strftime('%Y-%m-%d %H:%M')),
            ]
            
            for field, value in data:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            if analysis.result_image:
                doc.add_paragraph(f"Result image available: {analysis.result_image.name}")
            
            doc.add_paragraph()  # Spacing
    
    def _add_test_notes(self, doc, test_notes):
        """Add test notes"""
        doc.add_heading('Test Notes', level=1)
        
        for note in test_notes:
            test_name = f"Test {note.test_number:02d}"
            doc.add_heading(f'{test_name}', level=2)
            doc.add_paragraph(note.note_text)
            doc.add_paragraph(f"Created: {note.created_at.strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph()  # Spacing
    
    def _add_summary(self, doc, qa_record):
        """Add summary and conclusions"""
        doc.add_heading('Summary and Conclusions', level=1)
        
        # Count tests
        total_tests = 20
        completed_tests = sum(1 for i in range(1, 21) if getattr(qa_record, f'test_{i:02d}') is not None)
        passed_tests = 0
        failed_tests = 0
        
        # Check tolerance
        tolerances = {
            'test_01': 1.0, 'test_02': 1.0, 'test_03': 1.0, 'test_04': 1.0, 'test_05': 1.0,
            'test_06': 1.0, 'test_07': 1.0, 'test_08': 1.0, 'test_09': 1.0, 'test_10': 1.0,
            'test_11': 2.0, 'test_12': 1.0, 'test_13': 1.0, 'test_14': 1.0, 'test_15': 2.0,
            'test_16': 1.0, 'test_17': 3.0, 'test_18': 3.0, 'test_19': 1.0, 'test_20': 2.0,
        }
        
        for i in range(1, 21):
            field_name = f'test_{i:02d}'
            value = getattr(qa_record, field_name)
            tolerance = tolerances.get(field_name, 0)
            
            if value is not None:
                if abs(value) <= tolerance:
                    passed_tests += 1
                else:
                    failed_tests += 1
        
        # Summary table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        summary_data = [
            ('Total Tests', total_tests),
            ('Completed Tests', completed_tests),
            ('Passed Tests', passed_tests),
            ('Failed Tests', failed_tests),
            ('Completion Rate', f"{(completed_tests/total_tests)*100:.1f}%"),
            ('Pass Rate', f"{(passed_tests/completed_tests)*100:.1f}%" if completed_tests > 0 else "N/A"),
        ]
        
        for metric, value in summary_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)
        
        # Overall status
        doc.add_paragraph()
        if failed_tests == 0 and completed_tests > 0:
            status_para = doc.add_paragraph()
            status_para.add_run("Overall Status: PASS").bold = True
        elif failed_tests > 0:
            status_para = doc.add_paragraph()
            status_para.add_run(f"Overall Status: FAIL ({failed_tests} tests failed)").bold = True
        else:
            status_para = doc.add_paragraph()
            status_para.add_run("Overall Status: INCOMPLETE").bold = True
        
        # Notes
        if qa_record.notes:
            doc.add_paragraph()
            doc.add_heading('Additional Notes', level=2)
            doc.add_paragraph(qa_record.notes)
    
    def save_report(self, doc, filename, output_format='pdf'):
        """Save the document to a temporary file"""
        temp_dir = tempfile.gettempdir()
        temp_docx = os.path.join(temp_dir, f"{filename}.docx")
        temp_pdf = os.path.join(temp_dir, f"{filename}.pdf")
        
        # Save DOCX first
        doc.save(temp_docx)
        
        if output_format == 'docx':
            return temp_docx

        ok, errors = convert_docx_to_pdf(temp_docx, temp_pdf)
        if ok and os.path.exists(temp_pdf) and os.path.getsize(temp_pdf) > 0:
            return temp_pdf
        raise RuntimeError("PDF conversion failed. " + " | ".join(errors))
    
    def generate_report_response(self, qa_record, filename=None, output_format='pdf', preview=False):
        """Generate a Django HttpResponse with the PDF file"""
        if filename is None:
            filename = f"qa_report_{qa_record.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        output_format = (output_format or 'pdf').strip().lower()
        if output_format not in ('pdf', 'docx'):
            output_format = 'pdf'
        
        # Generate the document
        doc = self.generate_qa_report(qa_record)
        
        try:
            temp_file = self.save_report(doc, filename, output_format=output_format)
            if output_format == 'pdf' and not temp_file.endswith('.pdf'):
                raise RuntimeError("PDF export requested but non-PDF file was produced.")
        except Exception as exc:
            logger.exception("QA report export failed.")
            if output_format == 'pdf':
                return HttpResponse(
                    "PDF export failed after template-to-DOCX conversion. "
                    "Please verify Microsoft Word COM registration (or install LibreOffice) and try again. "
                    f"Technical detail: {exc}",
                    status=500,
                    content_type='text/plain; charset=utf-8',
                )
            return HttpResponse(
                f"Failed to export report: {exc}",
                status=500,
                content_type='text/plain; charset=utf-8',
            )
        
        # Determine content type and file extension
        if temp_file.endswith('.pdf'):
            content_type = 'application/pdf'
            file_extension = 'pdf'
        else:
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_extension = 'docx'
        
        # Read the file and create response
        with open(temp_file, 'rb') as f:
            response = HttpResponse(f.read(), content_type=content_type)
            disposition = 'inline' if preview else 'attachment'
            response['Content-Disposition'] = f'{disposition}; filename="{filename}.{file_extension}"'
        
        # Clean up temporary file
        try:
            os.remove(temp_file)
        except:
            pass
        
        return response


# ============================================================================
# View Function for Report Generation
# ============================================================================

def generate_qa_report_view(request, qa_record_id):
    """
    View function to generate QA report.
    
    This function is called via URL to generate and download QA reports.
    Returns a PDF or DOCX file response.
    """
    try:
        qa_record = QARecord.objects.get(id=qa_record_id)
        if qa_record.is_draft:
            return HttpResponse(
                "Cannot generate report for a draft QA record. Please resume and submit final QA first.",
                status=400,
                content_type='text/plain; charset=utf-8',
            )
        generator = QAReportGenerator()
        output_format = request.GET.get('format', 'pdf').lower()
        if output_format not in ('pdf', 'docx'):
            output_format = 'pdf'
        preview = request.GET.get('preview') in ('1', 'true', 'yes')
        return generator.generate_report_response(
            qa_record,
            output_format=output_format,
            preview=preview,
        )
    except QARecord.DoesNotExist:
        from django.http import HttpResponseNotFound
        return HttpResponseNotFound("QA record not found")
    except Exception as e:
        from django.http import HttpResponseServerError
        return HttpResponseServerError(f"Error generating report: {str(e)}") 